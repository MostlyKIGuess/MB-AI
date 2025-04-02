import streamlit as st
import os
from google import genai
from google.genai import types
from google.genai.types import GenerateContentConfig
from fuzzywuzzy import fuzz
from docx import Document

st.set_page_config(
    page_title="Music Blocks Assistant",
    page_icon="🎵",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    body {
        background: #2b2b2b; 
        font-family: 'Arial', sans-serif;
    }
    .stApp {
        max-width: 900px;
        margin: 0 auto;
        background-color: #2b2b2b;
        border-radius: 15px;
        padding: 20px;
        box-shadow: 0px 4px 20px rgba(0, 0, 0, 0.5);
        color: #ffffff;
    }
    .st-expander {
        background-color: #3a3a3a;
        border: 2px solid #4caf50;
        border-radius: 10px;
        padding: 10px;
    }
    .st-expander-header {
        font-size: 18px;
        font-weight: bold;
        color: #4caf50;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 1rem;
        display: flex;
        align-items: flex-start;
        font-family: 'Comic Sans MS', cursive, sans-serif;
    }
    .chat-message.user {
        background-color: #4caf50;
        color: #ffffff;
        border: 2px solid #388e3c;
    }
    .chat-message.bot {
        background-color: #1e88e5;
        color: #ffffff;
        border: 2px solid #1565c0;
    }
    .chat-message .avatar {
        width: 50px;
        height: 50px;
        border-radius: 50%;
        margin-right: 10px;
        background-color: #ffffff;
        display: flex;
        justify-content: center;
        align-items: center;
        font-size: 20px;
        font-weight: bold;
        color: #000000;
    }
    .chat-message .message {
        flex: 1;
        padding: 0 10px;
    }
    .stButton>button {
        background-color: #4caf50;
        color: white;
        border: none;
        padding: 10px 20px;
        font-size: 16px;
        border-radius: 5px;
        cursor: pointer;
        transition: background-color 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #388e3c;
    }
    .stTextInput>div>div>input {
        border: 2px solid #4caf50;
        border-radius: 5px;
        padding: 10px;
        font-size: 16px;
        background-color: #3a3a3a;
        color: #ffffff;
    }
</style>
""", unsafe_allow_html=True)

if 'GOOGLE_API_KEY' not in st.secrets:
    st.warning("Please set your Google API key in Streamlit secrets.")
    api_key = st.text_input("Enter your Google API Key:", type="password")
    if api_key:
        client = genai.Client(api_key=api_key)
else:
    client = genai.Client(api_key=st.secrets["GOOGLE_API_KEY"])

os.makedirs('lesson_plans', exist_ok=True)

APPRECIATION_PHRASES = ["thanks", "thank you", "great", "good job", "well done", "appreciate it", "good"]
DISCOURAGEMENT_PHRASES = ["not helpful", "bad", "poor", "wrong", "incorrect", "disappointed"]

APPRECIATION_RESPONSE = "You're welcome! I'm glad I could help. If you have more questions, feel free to ask."
DISCOURAGEMENT_RESPONSE = "I'm sorry to hear that you are not satisfied with my response. Could you please provide more details or ask another question so I can assist you better?"

def detect_sentiment(message):
    for phrase in APPRECIATION_PHRASES:
        if fuzz.ratio(message.lower(), phrase.lower()) > 90:
            return 'appreciation'
    for phrase in DISCOURAGEMENT_PHRASES:
        if fuzz.ratio(message.lower(), phrase.lower()) > 90:
            return 'discouragement'
    return None

greetings = ["hi", "hello", "hey", "greetings", "good morning", "good afternoon", "good evening"]

def is_greeting(user_input):
    user_input = user_input.strip().lower()
    for greeting in greetings:
        if fuzz.ratio(user_input, greeting) > 90:
            return True
    return False

def generate_response(query, is_first_message=False):
    if 'chat_session' not in st.session_state:
        st.session_state.chat_session = client.chats.create(
            model="gemini-2.0-flash",
            config=types.GenerateContentConfig(
                system_instruction="""You are Music Blocks Assistant, an expert in music education.

Your primary role is to help teachers and students work with Music Blocks, a programming environment that helps young people explore musical concepts through coding.

When responding:
- Provide clear, age-appropriate explanations about musical concepts
- Suggest engaging coding activities that connect programming with musical exploration
- Create structured lesson plans when requested, including learning objectives, materials needed, step-by-step activities, and assessment strategies
- Include examples of how to implement musical ideas 
- Be encouraging and foster a growth mindset toward both music and programming
- Reference the provided documentation to ensure accuracy in your responses
- Maintain memory of the entire conversation history, referring back to previous exchanges when relevant

Always maintain a supportive, enthusiastic tone that inspires creativity at the intersection of music and technology."""
            )
        )
        
        if is_first_message and hasattr(st.session_state, 'guide_context'):
            st.session_state.chat_session.send_message(
                f"Here is the Music Blocks documentation to reference when answering questions:\n\n{st.session_state.guide_context}\n\nPlease use this documentation to help answer user questions accurately."
            )
    
    response = st.session_state.chat_session.send_message(query)
    
    return response.text

def save_response_to_docx(response, topic):
    doc = Document()
    doc.add_heading('Music Blocks Lesson Plan', 0)
    doc.add_paragraph(response)
    filename = f"lesson_plan_{topic.replace(' ', '_')}.docx"
    save_path = os.path.join('lesson_plans', filename)
    doc.save(save_path)
    return save_path

if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'lesson_plan' not in st.session_state:
    st.session_state.lesson_plan = ""
if 'topic' not in st.session_state:
    st.session_state.topic = ""
if 'is_first_message' not in st.session_state:
    st.session_state.is_first_message = True

with open("guide.md", "r") as guide_file:
    guide_context = guide_file.read()
    st.session_state.guide_context = guide_context

with st.expander("🎵 About Music Blocks Assistant", expanded=True):
    st.markdown("""
    **Music Blocks Assistant** helps you:
    
    * Learn about music concepts
    * Generate music lesson plans
    * Explore musical topics
    
    Try asking for a lesson plan or any questions about music!
    """)

st.title("🎵 Music Blocks Assistant")
st.markdown("""
This is a Music Blocks lesson plan creation assistant. Ask questions about music or request a lesson plan!
""")

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

prompt = st.chat_input("Ask something about Music Blocks...")

if prompt:
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    with st.chat_message("user"):
        st.markdown(prompt)
    
    sentiment = detect_sentiment(prompt)
    if sentiment == 'appreciation':
        response = APPRECIATION_RESPONSE
    elif sentiment == 'discouragement':
        response = DISCOURAGEMENT_RESPONSE
    elif is_greeting(prompt):
        response = "Hello! How can I assist you today?"
    else:
        response = generate_response(prompt, st.session_state.is_first_message)
        if st.session_state.is_first_message:
            st.session_state.is_first_message = False
    
    with st.chat_message("assistant"):
        st.markdown(response)
    
    st.session_state.messages.append({"role": "assistant", "content": response})

