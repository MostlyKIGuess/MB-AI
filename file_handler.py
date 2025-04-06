import os
from docx import Document

def ensure_directories():
    """Ensure necessary directories exist."""
    os.makedirs('lesson_plans', exist_ok=True)

def load_guide_context():
    """Load the guide context from file."""
    try:
        with open("guide.md", "r") as guide_file:
            return guide_file.read()
    except FileNotFoundError:
        return "Guide file not found."

def save_response_to_docx(response, topic):
    """Save AI response to a Word document."""
    doc = Document()
    doc.add_heading('Music Blocks Lesson Plan', 0)
    doc.add_paragraph(response)
    filename = f"lesson_plan_{topic.replace(' ', '_')}.docx"
    save_path = os.path.join('lesson_plans', filename)
    doc.save(save_path)
    return save_path
